import docx
#from data import import_dictionary


def doc(dict,file):
    document = docx.Document(file)
    for p in document.paragraphs:
        inline = p.runs
        for i in range(len(inline)):
            text = inline[i].text
            for key in dict.keys():
                if key in text:
                     text=text.replace(key,dict[key])
                     inline[i].text = text

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    inline = p.runs
                    for i in range(len(inline)):
                        text = inline[i].text
                        for key in dict.keys():
                            if key in text:
                                text=text.replace(key,dict[key])
                                inline[i].text = text

    document.save('new.docx')

#doc(import_dictionary)

